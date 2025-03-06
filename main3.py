import os
import json
import datetime
import tkinter as tk
from tkinter import ttk, messagebox
import pythoncom
import comtypes.client
import docx
from docx import Document
import sys

# ---------- Пути к JSON и шаблонам ----------
EMPLOYEES_JSON = "employees.json"
CONFIG_JSON = "config.json"

# TEMPLATES_FOLDER = os.path.join(os.getcwd(), "Templates")

# Новое
EXECUTABLE_DIR = os.path.dirname(sys.executable)
TEMPLATES_FOLDER = os.path.join(EXECUTABLE_DIR, "Templates")

TEMPLATE_OBLOZHKA        = os.path.join(TEMPLATES_FOLDER, "oblozhka.docx")
TEMPLATE_PRODLENIE       = os.path.join(TEMPLATES_FOLDER, "prodlenie.docx")
TEMPLATE_POSTANOVLENIE   = os.path.join(TEMPLATES_FOLDER, "postanovlenie.docx")
TEMPLATE_SOPROVODITELNOE = os.path.join(TEMPLATES_FOLDER, "soprovoditelnoe.docx")

# Папки для сохранения на рабочем столе
DESKTOP = os.path.join(os.path.join(os.environ["USERPROFILE"]), "Desktop")
CASCO_DOCS_FOLDER = os.path.join(DESKTOP, "CASCO DOCS")
CASCO_PDF_FOLDER = os.path.join(DESKTOP, "CASCO PDF")


# Путь к иконке
ICON_PATH = os.path.join(EXECUTABLE_DIR, "icon.ico")

def load_json(filepath):
    """Загружает JSON из файла. Если нет файла, возвращает None."""
    if not os.path.exists(filepath):
        return None
    with open(filepath, "r", encoding="utf-8") as f:
        return json.load(f)

def save_json(data, filepath):
    """Сохраняет данные в JSON-файл."""
    with open(filepath, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)

def ensure_output_folders():
    """Создаёт папки CASCO DOCS и CASCO PDF на рабочем столе, если они не существуют."""
    if not os.path.exists(CASCO_DOCS_FOLDER):
        os.mkdir(CASCO_DOCS_FOLDER)
    if not os.path.exists(CASCO_PDF_FOLDER):
        os.mkdir(CASCO_PDF_FOLDER)

# ---------------------------------------------------------
# Эвристика: определяем пол (male/female/unknown)
# ---------------------------------------------------------
def guess_gender(surname: str) -> str:
    """
    Простая эвристика для определения пола по фамилии.
    Возвращает 'male', 'female' или 'unknown'.
    """
    s = surname.strip().lower()

    female_endings = (
        "енская", "инская", "анская", "онская", "ицкая", "цкая", "ская",
        "ова", "ева", "ёва", "ина", "ына",
        "ая", "яя",
        "овна", "евна"
    )
    male_endings = (
        "енский", "инский", "анский", "онский", "цкий", "ский", "ченко",
        "нов", "ков", "ёв", "ев", "ин", "ын",
        "ий", "ый", "ой",
        "ук", "юк", "як",
        "ич",
        "арь", "яр",
        "енко"
    )

    # Проверяем женские
    for ending in sorted(female_endings, key=len, reverse=True):
        if s.endswith(ending):
            return "female"

    # Проверяем мужские
    for ending in sorted(male_endings, key=len, reverse=True):
        if s.endswith(ending):
            return "male"

    return "unknown"

# ---------------------------------------------------------
# Склоняем фамилию в дательный падеж (эвристика)
# ---------------------------------------------------------
def decline_dative(surname: str) -> str:
    """
    Склоняет фамилию в дательный падеж на основе эвристических правил:
    - Если пол = male: подставляем "ов"->"ову", "ев"->"еву" и т.д.
    - Если пол = female: "ова"->"овой", "ая"->"ой" и т.п.
    - Если unknown, возвращаем исходную фамилию.
    """
    gender = guess_gender(surname)
    s = surname.strip()

    if gender == "male":
        # мужские фамилии
        if s.endswith("ов"):
            return s[:-2] + "ову"
        elif s.endswith("ев"):
            return s[:-2] + "еву"
        elif s.endswith("ин"):
            return s[:-2] + "ину"
        elif s.endswith("ын"):
            return s[:-2] + "ыну"
        elif s.endswith("ский"):
            return s[:-4] + "скому"
        elif s.endswith("цкий"):
            return s[:-4] + "цкому"
        elif s.endswith("ий"):
            return s[:-2] + "ию"
        elif s.endswith("ый"):
            return s[:-2] + "ому"
        else:
            # по умолчанию для мужской: добавляем "у"
            return s + "у"
    elif gender == "female":
        # женские фамилии
        if s.endswith(("ова", "ева", "ёва", "ина", "ына", "овна", "евна")):
            # Иванова -> Ивановой
            return s[:-1] + "ой"
        elif s.endswith(("ая", "яя")):
            # Большая -> Большой
            return s[:-2] + "ой"
        elif s.endswith(("ская", "цкая")):
            # ская -> ской, цкая -> цкой
            return s[:-2] + "ой"
        else:
            # по умолчанию для женской: добавляем "ой"
            return s + "ой"
    else:
        # unknown
        return s

def replace_placeholders_in_docx(template_path, placeholders, output_path):
    """
    Открывает шаблон Word (docx),
    заменяет все вхождения {ключ} -> значение по run'ам (сохраняет формат),
    а затем сохраняет в output_path.
    """
    doc_ = docx.Document(template_path)

    # Замена в параграфах
    for p in doc_.paragraphs:
        for run in p.runs:
            for key, val in placeholders.items():
                if key in run.text:
                    run.text = run.text.replace(key, val)

    # Замена в таблицах
    for table in doc_.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        for key, val in placeholders.items():
                            if key in run.text:
                                run.text = run.text.replace(key, val)

    doc_.save(output_path)



def open_file(path):
    """Открывает файл (PDF или другое) в системе по умолчанию."""
    os.startfile(path)

# -------------- Окно редактирования сотрудников --------------
class EmployeesEditor(tk.Toplevel):
    def __init__(self, master, employees):
        super().__init__(master)
        self.title("Редактирование сотрудников")
        self.geometry("600x300")
        self.employees = employees

        columns = ("Фамилия", "Инициалы", "Звание", "Должность", "Телефон")
        self.tree = ttk.Treeview(self, columns=columns, show='headings')
        for col in columns:
            self.tree.heading(col, text=col)
            self.tree.column(col, width=100)
        self.tree.pack(fill=tk.BOTH, expand=True)

        frame_btn = tk.Frame(self)
        frame_btn.pack(fill=tk.X)

        btn_add = tk.Button(frame_btn, text="Добавить", command=self.add_employee)
        btn_add.pack(side=tk.LEFT, padx=5, pady=5)

        btn_edit = tk.Button(frame_btn, text="Редактировать", command=self.edit_employee)
        btn_edit.pack(side=tk.LEFT, padx=5, pady=5)

        btn_delete = tk.Button(frame_btn, text="Удалить", command=self.delete_employee)
        btn_delete.pack(side=tk.LEFT, padx=5, pady=5)

        self.load_employees()

    def load_employees(self):
        for row in self.tree.get_children():
            self.tree.delete(row)
        for emp in self.employees:
            self.tree.insert("", tk.END, values=(
                emp.get("Фамилия", ""),
                emp.get("Инициалы", ""),
                emp.get("Звание", ""),
                emp.get("Должность", ""),
                emp.get("Телефон", "")
            ))

    def add_employee(self):
        EmployeeForm(self, self.employees, mode="add", on_save=self.on_employee_saved)

    def edit_employee(self):
        selected = self.tree.focus()
        if not selected:
            messagebox.showwarning("Внимание", "Выберите сотрудника из списка")
            return
        index = self.tree.index(selected)
        EmployeeForm(self, self.employees, mode="edit", on_save=self.on_employee_saved, index=index)

    def delete_employee(self):
        selected = self.tree.focus()
        if not selected:
            messagebox.showwarning("Внимание", "Выберите сотрудника из списка")
            return
        index = self.tree.index(selected)
        confirm = messagebox.askyesno("Подтверждение", "Удалить выбранного сотрудника?")
        if confirm:
            self.employees.pop(index)
            self.load_employees()

    def on_employee_saved(self):
        self.load_employees()

class EmployeeForm(tk.Toplevel):
    def __init__(self, master, employees, mode="add", on_save=None, index=None):
        super().__init__(master)
        self.title("Сотрудник")
        self.geometry("300x250")
        self.employees = employees
        self.on_save = on_save
        self.mode = mode
        self.index = index

        tk.Label(self, text="Фамилия").pack()
        self.entry_lastname = tk.Entry(self)
        self.entry_lastname.pack()

        tk.Label(self, text="Инициалы").pack()
        self.entry_initials = tk.Entry(self)
        self.entry_initials.pack()

        tk.Label(self, text="Звание").pack()
        self.entry_rank = tk.Entry(self)
        self.entry_rank.pack()

        tk.Label(self, text="Должность").pack()
        self.entry_position = tk.Entry(self)
        self.entry_position.pack()

        tk.Label(self, text="Телефон").pack()
        self.entry_phone = tk.Entry(self)
        self.entry_phone.pack()

        if self.mode == "edit" and self.index is not None:
            emp = self.employees[self.index]
            self.entry_lastname.insert(0, emp.get("Фамилия", ""))
            self.entry_initials.insert(0, emp.get("Инициалы", ""))
            self.entry_rank.insert(0, emp.get("Звание", ""))
            self.entry_position.insert(0, emp.get("Должность", ""))
            self.entry_phone.insert(0, emp.get("Телефон", ""))

        tk.Button(self, text="Сохранить", command=self.save_employee).pack(pady=10)

    def save_employee(self):
        new_emp = {
            "Фамилия": self.entry_lastname.get().strip(),
            "Инициалы": self.entry_initials.get().strip(),
            "Звание": self.entry_rank.get().strip(),
            "Должность": self.entry_position.get().strip(),
            "Телефон": self.entry_phone.get().strip()
        }
        if self.mode == "add":
            self.employees.append(new_emp)
        else:
            self.employees[self.index] = new_emp
        if self.on_save:
            self.on_save()
        self.destroy()


import pythoncom
import comtypes.client


class App(tk.Tk):
    def __init__(self):
        super().__init__()
        self.title("КАСКО - Автоматизация документов        15 отдел полиции           v. 1.0.1")
        self.geometry("900x600")  # Ширина и высота окна

        # Устанавливаем иконку
        if os.path.exists(ICON_PATH):
            self.iconbitmap(ICON_PATH)
        else:
            print(f"⚠️ Иконка не найдена: {ICON_PATH}")

        self.temp_folder = os.path.join(os.path.dirname(__file__), "Temp")
        if not os.path.exists(self.temp_folder):
            os.mkdir(self.temp_folder)



        ensure_output_folders()


        # Инициализация COM
        pythoncom.CoInitialize()

        # Открываем Word ОДИН раз
        self.word_app = comtypes.client.CreateObject("Word.Application")
        self.word_app.Visible = False



        # Загрузка сотрудников
        self.employees = load_json(EMPLOYEES_JSON)
        if self.employees is None:
            self.employees = []

        # Загрузка config.json (сохранённый индекс сотрудника)
        self.config_data = load_json(CONFIG_JSON)
        if self.config_data is None:
            self.config_data = {}

        # Создаём интерфейс
        self.create_widgets()

        # Восстанавливаем настройки (выбор сотрудника в combo)
        self.load_config()

        self.protocol("WM_DELETE_WINDOW", self.on_close)

    def create_widgets(self):
        """Создаёт все элементы интерфейса."""
        frame_fields = tk.Frame(self)
        frame_fields.pack(fill=tk.X, padx=10, pady=10)

        # KUSP
        tk.Label(frame_fields, text="КУСП:", width=20, anchor="e").grid(row=0, column=0, padx=5, pady=5, sticky="e")
        self.var_kusp = tk.StringVar()
        tk.Entry(frame_fields, textvariable=self.var_kusp, width=20).grid(row=0, column=1, padx=5, pady=5, sticky="w")

        # Дата регистрации КУСП
        tk.Label(frame_fields, text="Дата регистрации КУСП:", anchor="e").grid(row=1, column=0, padx=5, pady=5, sticky="e")
        self.var_date_reg = tk.StringVar()
        tk.Entry(frame_fields, textvariable=self.var_date_reg, width=15).grid(row=1, column=1, padx=5, pady=5, sticky="w")

        # Дата постановления
        tk.Label(frame_fields, text="Дата постановления:", anchor="e").grid(row=2, column=0, padx=5, pady=5, sticky="e")
        self.var_order_date = tk.StringVar()
        tk.Entry(frame_fields, textvariable=self.var_order_date, width=15).grid(row=2, column=1, padx=5, pady=5, sticky="w")

        # Фамилия заявителя
        tk.Label(frame_fields, text="Фамилия заявителя:", anchor="e").grid(row=3, column=0, padx=5, pady=5, sticky="e")
        self.var_last_name = tk.StringVar()
        tk.Entry(frame_fields, textvariable=self.var_last_name, width=30).grid(row=3, column=1, padx=5, pady=5, sticky="w")

        # Инициалы заявителя
        tk.Label(frame_fields, text="Инициалы заявителя:", anchor="e").grid(row=4, column=0, padx=5, pady=5, sticky="e")
        self.var_initials = tk.StringVar()
        tk.Entry(frame_fields, textvariable=self.var_initials, width=30).grid(row=4, column=1, padx=5, pady=5, sticky="w")

        # Адрес проживания
        tk.Label(frame_fields, text="Адрес проживания:", anchor="e").grid(row=5, column=0, padx=5, pady=5, sticky="e")
        self.var_address_applicant = tk.StringVar()
        tk.Entry(frame_fields, textvariable=self.var_address_applicant, width=53).grid(row=5, column=1, padx=5, pady=5, sticky="w")

        # Марка и модель авто
        tk.Label(frame_fields, text="Марка и модель авто:", anchor="e").grid(row=6, column=0, padx=5, pady=5, sticky="e")
        self.var_car_brand = tk.StringVar()
        tk.Entry(frame_fields, textvariable=self.var_car_brand, width=53).grid(row=6, column=1, padx=5, pady=5, sticky="w")

        # Госномер
        tk.Label(frame_fields, text="Госномер автомобиля:", anchor="e").grid(row=7, column=0, padx=5, pady=5, sticky="e")
        self.var_car_plate = tk.StringVar()
        tk.Entry(frame_fields, textvariable=self.var_car_plate, width=20).grid(row=7, column=1, padx=5, pady=5, sticky="w")

        # Дата обнаружения
        tk.Label(frame_fields, text="Дата обнаружения:", anchor="e").grid(row=8, column=0, padx=5, pady=5, sticky="e")
        self.var_damage_date = tk.StringVar()
        tk.Entry(frame_fields, textvariable=self.var_damage_date, width=15).grid(row=8, column=1, padx=5, pady=5, sticky="w")

        # Время обнаружения
        tk.Label(frame_fields, text="Время обнаружения:", anchor="e").grid(row=9, column=0, padx=5, pady=5, sticky="e")
        self.var_damage_time = tk.StringVar()
        tk.Entry(frame_fields, textvariable=self.var_damage_time, width=8).grid(row=9, column=1, padx=5, pady=5, sticky="w")

        # Повреждения обнаружил у дома + «в Санкт-Петербурге»
        tk.Label(frame_fields, text="Повреждения обнаружил у дома:", anchor="e").grid(row=10, column=0, padx=5, pady=5, sticky="e")
        discovery_frame = tk.Frame(frame_fields)
        discovery_frame.grid(row=10, column=1, padx=5, pady=5, sticky="w")

        self.var_discovery_address = tk.StringVar()
        tk.Entry(discovery_frame, textvariable=self.var_discovery_address, width=30).pack(side=tk.LEFT)
        tk.Label(discovery_frame, text="в Санкт-Петербурге").pack(side=tk.LEFT, padx=5)

        # Перечень повреждений (многострочное поле)
        tk.Label(frame_fields, text="Перечень повреждений:", anchor="e").grid(row=11, column=0, padx=5, pady=5, sticky="ne")
        self.txt_damages = tk.Text(frame_fields, width=40, height=3)
        self.txt_damages.grid(row=11, column=1, padx=5, pady=5, sticky="w")
        # Чтобы по Tab фокус уходил дальше
        self.txt_damages.bind("<Tab>", self.focus_next_widget)

        # ----- Выбор сотрудника -----
        tk.Label(frame_fields, text="Сотрудник:", anchor="e").grid(row=0, column=2, padx=5, pady=5, sticky="e")
        self.combo_employee = ttk.Combobox(frame_fields, state="readonly")
        self.combo_employee.grid(row=0, column=3, padx=5, pady=5, sticky="w")

        btn_edit_employees = tk.Button(frame_fields, text="Изменить список", command=self.on_edit_employees)
        btn_edit_employees.grid(row=0, column=4, padx=5, pady=5, sticky="w")

        # Кнопка «Очистить»
        self.btn_clear = tk.Button(self, text="Очистить", command=self.clear_fields)
        self.btn_clear.pack(pady=5)

        # Панель кнопок (снизу)
        frame_buttons = tk.Frame(self)
        frame_buttons.pack(side=tk.BOTTOM, fill=tk.X, pady=10)

        btn_open_oblozhka = tk.Button(frame_buttons, text="Открыть обложку", command=self.open_oblozhka)
        btn_open_oblozhka.pack(side=tk.LEFT, padx=5)

        btn_open_prodlenie = tk.Button(frame_buttons, text="Открыть продление", command=self.open_prodlenie)
        btn_open_prodlenie.pack(side=tk.LEFT, padx=5)

        btn_save_postanovlenie = tk.Button(
            frame_buttons,
            text="Сохранить и открыть постановление",
            command=self.save_and_open_postanovlenie
        )
        btn_save_postanovlenie.pack(side=tk.LEFT, padx=5)

        btn_save_soprov = tk.Button(
            frame_buttons,
            text="Сохранить и открыть сопровод",
            command=self.save_and_open_soprovod
        )
        btn_save_soprov.pack(side=tk.LEFT, padx=5)



    def export_to_pdf(self, docx_path, pdf_path):
        """
        Экспортирует docx_path -> pdf_path, используя 
        один экземпляр Word (self.word_app).
        """
        doc = self.word_app.Documents.Open(docx_path)
        doc.ExportAsFixedFormat(
            pdf_path,
            ExportFormat=17,  # wdExportFormatPDF
            OpenAfterExport=False,
            OptimizeFor=0,
            CreateBookmarks=1,
            DocStructureTags=True,
            BitmapMissingFonts=True,
            UseISO19005_1=False
        )
        doc.Close(False)











    def focus_next_widget(self, event):
        """Чтобы Tab в Text переключал фокус, а не вставлял табуляцию."""
        event.widget.tk_focusNext().focus_set()
        return "break"

    def on_edit_employees(self):
        editor = EmployeesEditor(self, self.employees)
        self.wait_window(editor)
        self.refresh_employee_list()

    def clear_fields(self):
        """Очистить все поля."""
        self.var_kusp.set("")
        self.var_date_reg.set("")
        self.var_order_date.set("")
        self.var_last_name.set("")
        self.var_initials.set("")
        self.var_address_applicant.set("")
        self.var_car_brand.set("")
        self.var_car_plate.set("")
        self.var_damage_date.set("")
        self.var_damage_time.set("")
        self.var_discovery_address.set("")
        self.txt_damages.delete("1.0", tk.END)

    def refresh_employee_list(self):
        """Обновляет список сотрудников в combo_employee."""
        emp_names = []
        for emp in self.employees:
            ln = emp.get("Фамилия", "")
            ini = emp.get("Инициалы", "")
            emp_names.append(f"{ln} {ini}")
        self.combo_employee["values"] = emp_names

    def load_config(self):
        """Восстанавливаем выбранного сотрудника из config."""
        self.refresh_employee_list()

        # Берём индекс из config
        selected_emp_index = self.config_data.get("selected_employee_index", 0)
        
        # Если он < 0, ставим 0
        if selected_emp_index < 0:
            selected_emp_index = 0

        # Если он >= длины списка, ставим 0 (или ставим последнего)
        if selected_emp_index >= len(self.employees) and len(self.employees) > 0:
            selected_emp_index = 0

        # Если employees не пуст, устанавливаем current
        if len(self.employees) > 0:
            self.combo_employee.current(selected_emp_index)


    def on_close(self):
        # 1. Сохраняем настройки (config) и список сотрудников (employees)
        idx = self.combo_employee.current()
        self.config_data["selected_employee_index"] = idx
        save_json(self.config_data, CONFIG_JSON)
        save_json(self.employees, EMPLOYEES_JSON)

        # 2. Закрываем Word, если он запущен
        if getattr(self, 'word_app', None) is not None:
            self.word_app.Quit()
            self.word_app = None

        # 3. Деинициализируем COM
        pythoncom.CoUninitialize()

        # 4. Удаляем все файлы из временной папки (Temp)
        if os.path.exists(self.temp_folder):
            for fname in os.listdir(self.temp_folder):
                file_path = os.path.join(self.temp_folder, fname)
                try:
                    os.remove(file_path)
                except Exception as e:
                    print(f"Не удалось удалить {file_path}: {e}")
            # Если хотите удалить и саму папку (если она пуста):
            try:
                os.rmdir(self.temp_folder)
            except OSError:
                pass  # Если папка не пуста или что-то пошло не так

        # 5. Закрываем главное окно
        self.destroy()



    def get_selected_employee(self):
        """
        Возвращает словарь сотрудника, выбранного в combo_employee,
        или None, если индекс некорректный.
        """
        idx = self.combo_employee.current()
        if idx < 0 or idx >= len(self.employees):
            return None
        return self.employees[idx]

    # ================= ВАЖНО: Эвристика дательного падежа =================
    def make_placeholders(self):
        """
        Формируем словарь плейсхолдеров для всех общих полей.
        Добавляем {LAST_NAME_DAT} - дательный падеж фамилии.
        """
        emp = self.get_selected_employee() or {}

        # Исходная фамилия (именительный)
        family_nom = self.var_last_name.get().strip()
        # Склоняем по нашим правилам:
        family_dat = decline_dative(family_nom)

        placeholders = {
            "{KUSP}": self.var_kusp.get(),
            "{DATE_REGISTRATION}": self.var_date_reg.get(),
            "{ORDER_DATE}": self.var_order_date.get(),
            # Именительный
            "{LAST_NAME}": family_nom,
            # Дательный
            "{LAST_NAME_DAT}": family_dat,

            "{INITIALS}": self.var_initials.get(),
            "{APPLICANT_ADDRESS}": self.var_address_applicant.get(),
            "{CAR_BRAND}": self.var_car_brand.get(),
            "{CAR_PLATE}": self.var_car_plate.get(),
            "{DAMAGE_DATE}": self.var_damage_date.get(),
            "{DAMAGE_TIME}": self.var_damage_time.get(),
            "{DISCOVERY_ADDRESS}": self.var_discovery_address.get(),
            "{DAMAGES}": self.txt_damages.get("1.0", tk.END).strip(),

            "{EMP_LASTNAME}": emp.get("Фамилия", ""),
            "{EMP_INITIALS}": emp.get("Инициалы", ""),
            "{EMP_RANK}": emp.get("Звание", ""),
            "{EMP_POSITION}": emp.get("Должность", ""),
            "{EMP_PHONE}": emp.get("Телефон", "")
        }
        return placeholders

    def get_date_plus_2(self):
        """
        Берём дату регистрации (var_date_reg),
        парсим как dd.mm.yyyy и прибавляем 2 дня.
        Возвращаем строку в том же формате.
        Если парсинг не удался, возвращаем пустую строку.
        """
        raw_date = self.var_date_reg.get().strip()
        try:
            dt = datetime.datetime.strptime(raw_date, "%d.%m.%Y")
            dt_plus_2 = dt + datetime.timedelta(days=2)
            return dt_plus_2.strftime("%d.%m.%Y")
        except ValueError:
            return ""

    # ---------------------- Кнопки документов ----------------------
    def open_oblozhka(self):
        """
        Открыть обложку (PDF) — сохраняем временные файлы в папку 'Temp'.
        """
        if not os.path.exists(TEMPLATE_OBLOZHKA):
            messagebox.showerror("Ошибка", f"Не найден шаблон обложки: {TEMPLATE_OBLOZHKA}")
            return

        placeholders = self.make_placeholders()
        now_str = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")

        temp_docx = os.path.join(self.temp_folder, f"temp_oblozhka_{now_str}.docx")
        temp_pdf  = os.path.join(self.temp_folder, f"temp_oblozhka_{now_str}.pdf")

        replace_placeholders_in_docx(TEMPLATE_OBLOZHKA, placeholders, temp_docx)
        try:
            self.export_to_pdf(temp_docx, temp_pdf)
            open_file(temp_pdf)
        except Exception as e:
            messagebox.showerror("Ошибка", str(e))
        # НЕ удаляем в конце, чтобы пользователь мог распечатать
        # Но удалим всё скопом при выходе из программы (on_close).



    def open_prodlenie(self):
        """
        Открыть продление (PDF) — сохраняем временные файлы в 'Temp'.
        """
        if not os.path.exists(TEMPLATE_PRODLENIE):
            messagebox.showerror("Ошибка", f"Не найден шаблон продления: {TEMPLATE_PRODLENIE}")
            return

        placeholders = self.make_placeholders()
        placeholders["{ORDER_DATE_PRODLENIE}"] = self.get_date_plus_2()

        now_str = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        temp_docx = os.path.join(self.temp_folder, f"temp_prodlenie_{now_str}.docx")
        temp_pdf  = os.path.join(self.temp_folder, f"temp_prodlenie_{now_str}.pdf")

        replace_placeholders_in_docx(TEMPLATE_PRODLENIE, placeholders, temp_docx)
        try:
            self.export_to_pdf(temp_docx, temp_pdf)
            open_file(temp_pdf)
        except Exception as e:
            messagebox.showerror("Ошибка", str(e))
        # Не удаляем, user может напечатать.
        # Удалим потом при on_close.



    def save_and_open_postanovlenie(self):
        """
        Сохранить и открыть постановление (DOCX + PDF).
        DOCX -> CASCO DOCS с названием "{KUSP} от {DATE_REGISTRATION} (пост).docx"
        PDF  -> CASCO PDF  с названием "{KUSP} от {DATE_REGISTRATION} (пост).pdf"
        """
        if not os.path.exists(TEMPLATE_POSTANOVLENIE):
            messagebox.showerror("Ошибка", f"Не найден шаблон постановления: {TEMPLATE_POSTANOVLENIE}")
            return

        placeholders = self.make_placeholders()

        kusp_value   = placeholders["{KUSP}"]
        date_reg_val = placeholders["{DATE_REGISTRATION}"]

        # Формируем имена файлов
        docx_name = f"КУСП-{kusp_value} от {date_reg_val} (пост).docx"
        pdf_name  = f"КУСП-{kusp_value} от {date_reg_val} (пост).pdf"

        out_docx_path = os.path.join(CASCO_DOCS_FOLDER, docx_name)
        out_pdf_path  = os.path.join(CASCO_PDF_FOLDER,  pdf_name)

        replace_placeholders_in_docx(TEMPLATE_POSTANOVLENIE, placeholders, out_docx_path)
        try:
            self.export_to_pdf(out_docx_path, out_pdf_path)
            open_file(out_pdf_path)
        except Exception as e:
            messagebox.showerror("Ошибка", str(e))


    def save_and_open_soprovod(self):
        """
        Сохранить и открыть сопровод (только PDF).
        Название PDF: "{KUSP} от {DATE_REGISTRATION} (сопр).pdf"
        В CASCO PDF.
        Временный DOCX в CASCO DOCS удаляем потом.
        """
        if not os.path.exists(TEMPLATE_SOPROVODITELNOE):
            messagebox.showerror("Ошибка", f"Не найден шаблон сопроводительного: {TEMPLATE_SOPROVODITELNOE}")
            return

        placeholders = self.make_placeholders()

        kusp_value   = placeholders["{KUSP}"]
        date_reg_val = placeholders["{DATE_REGISTRATION}"]
        pdf_name     = f"КУСП-{kusp_value} от {date_reg_val} (сопр).pdf"

        temp_docx    = os.path.join(CASCO_DOCS_FOLDER, "temp_soprov.docx")
        out_pdf_path = os.path.join(CASCO_PDF_FOLDER, pdf_name)

        replace_placeholders_in_docx(TEMPLATE_SOPROVODITELNOE, placeholders, temp_docx)
        try:
            self.export_to_pdf(temp_docx, out_pdf_path)
            open_file(out_pdf_path)
        except Exception as e:
            messagebox.showerror("Ошибка", str(e))
        finally:
            if os.path.exists(temp_docx):
                os.remove(temp_docx)


def main():
    app = App()
    app.mainloop()

if __name__ == "__main__":
    main()
